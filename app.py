import streamlit as st
from streamlit_drawable_canvas import st_canvas
from pdf2image import convert_from_bytes
from PIL import Image, ImageDraw
import pytesseract
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from pdf2docx import Converter
from docx2pdf import convert
import io
import tempfile
import os

# =====================================================
# CONFIG
# =====================================================
st.set_page_config("Advanced PDF Editor", layout="wide")
st.title("📄 Advanced Professional PDF Editor")

MAX_MB = 15
MAX_PAGES = 15
DPI = 300
CANVAS_WIDTH = 900
OCR_LANG = "eng+hin"  # Multi-language OCR

# =====================================================
# UTILITIES
# =====================================================
def check_file(file):
    if file.size > MAX_MB * 1024 * 1024:
        st.error("❌ File too large")
        st.stop()

def pdf_to_images(file):
    file.seek(0)
    try:
        pages = convert_from_bytes(file.read(), dpi=DPI)
        if len(pages) > MAX_PAGES:
            st.error("❌ Too many pages")
            st.stop()
        return pages
    except Exception as e:
        st.error("PDF load failed (Poppler missing?)")
        st.error(e)
        st.stop()

def resize(img):
    ratio = CANVAS_WIDTH / img.width
    return img.resize((CANVAS_WIDTH, int(img.height * ratio)))

def save_pdf(images):
    if not images:
        st.error("No pages left to save")
        st.stop()
    buf = io.BytesIO()
    images[0].save(buf, format="PDF", save_all=True, append_images=images[1:])
    return buf.getvalue()

def apply_password(pdf_bytes, password):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    writer.encrypt(password)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

def add_watermark(img, text="CONFIDENTIAL"):
    wm = img.copy().convert("RGBA")
    overlay = Image.new("RGBA", wm.size, (255,255,255,0))
    draw = ImageDraw.Draw(overlay)
    draw.text((50, 50), text, fill=(255,0,0,120))
    return Image.alpha_composite(wm, overlay).convert("RGB")

def run_ocr(img):
    custom_config = r'--oem 3 --psm 6'
    return pytesseract.image_to_string(img, lang=OCR_LANG, config=custom_config)

# =====================================================
# SIDEBAR
# =====================================================
st.sidebar.markdown("## 🧰 PDF TOOLBOX")
st.sidebar.markdown("---")

tool = st.sidebar.radio(
    "Mode",
    ["Editor", "OCR", "PDF → Word (Editable)", "Word → PDF","Image → PDF", "Security"]
)

uploaded = st.sidebar.file_uploader("Upload PDF", type=["pdf"])
if not uploaded:
    st.info("Upload a PDF to start")
    st.stop()

check_file(uploaded)
pages = pdf_to_images(uploaded)

# =====================================================
# PAGE SELECT + DELETE + REORDER
# =====================================================
st.sidebar.markdown("### 📑 Pages")
cols = st.sidebar.columns(3)
selected_page = 0

for i, img in enumerate(pages):
    with cols[i % 3]:
        if st.button(f"{i+1}"):
            selected_page = i

if st.sidebar.button("🗑️ Delete Selected Page"):
    pages.pop(selected_page)
    st.experimental_rerun()

if len(pages) > 1:
    new_index = st.sidebar.slider("Move page to position", 1, len(pages), selected_page + 1) - 1
    if new_index != selected_page:
        pages.insert(new_index, pages.pop(selected_page))
        st.experimental_rerun()

page_img = pages[selected_page].convert("RGB")
canvas_bg = resize(page_img)

# =====================================================
# EDITOR MODE
# =====================================================
if tool == "Editor":
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        mode = st.selectbox("Tool", ["freedraw", "rect", "text", "pdf(editable)","signature"])
    with c2:
        size = st.slider("Size", 1, 40, 5)
    with c3:
        color = st.color_picker("Color", "#FF0000")
    with c4:
        rotate = st.selectbox("Rotate", [0, 90, 180, 270])

    if rotate:
        page_img = page_img.rotate(-rotate, expand=True)
        canvas_bg = resize(page_img)

    canvas = st_canvas(
        background_image=canvas_bg,
        drawing_mode=mode,
        stroke_width=size,
        stroke_color=color,
        height=canvas_bg.height,
        width=canvas_bg.width,
        key=f"canvas_{selected_page}"
    )

    if st.button("💾 Save Changes") and canvas.image_data is not None:
        edited = Image.fromarray(canvas.image_data.astype("uint8")).resize(page_img.size)
        final = page_img.convert("RGBA")
        final.alpha_composite(edited.convert("RGBA"))
        pages[selected_page] = final.convert("RGB")
        st.success("Saved")

    if st.button("💧 Add Watermark"):
        pages[selected_page] = add_watermark(page_img)
        st.success("Watermark added")

# =====================================================
# OCR MODE
# =====================================================
elif tool == "OCR":
    st.warning("OCR makes PDF searchable, not editable")

    if st.button("Run OCR on this page"):
        text = run_ocr(page_img)
        st.text_area("Extracted Text", text, height=300)

    if st.button("📄 OCR → Word (Editable)"):
        with st.spinner("Creating Word file..."):
            doc = Document()
            for i, img in enumerate(pages):
                text = run_ocr(img)
                doc.add_heading(f"Page {i+1}", level=2)
                doc.add_paragraph(text)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                with open(tmp.name, "rb") as f:
                    st.download_button(
                        "📥 Download OCR Word",
                        f.read(),
                        "ocr_editable.docx"
                    )
            os.remove(tmp.name)

# =====================================================
# PDF → Word
# =====================================================
elif tool == "PDF → Word (Editable)":
    if st.button("📝 Convert PDF to Word"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(uploaded.getbuffer())
            pdf_path = tmp_pdf.name

        docx_path = pdf_path.replace(".pdf", ".docx")

        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()

            with open(docx_path, "rb") as f:
                st.download_button(
                    "📥 Download Editable Word",
                    f.read(),
                    "editable.docx"
                )
        finally:
            os.remove(pdf_path)
            os.remove(docx_path)

# =====================================================
# TOOL: WORD -> PDF (FIXED FOR CLOUD)
# =====================================================
elif tool == "Word → PDF":
    st.header("📝 Convert Word to PDF")
    
    w_file = st.file_uploader("Upload Word File", type=["docx"])
    
    if w_file and st.button("Convert to PDF"):
        try:
            with st.spinner("Converting..."):
                # 1. Word file read karein
                doc = Document(w_file)
                
                # 2. PDF banayein
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # 3. Har paragraph ko PDF me likhein
                for p in doc.paragraphs:
                    # Encoding fix (Special characters handle karne ke liye)
                    safe_text = p.text.encode('latin-1', 'replace').decode('latin-1')
                    if safe_text.strip(): # Khali line ignore karein
                        pdf.multi_cell(0, 10, safe_text)
                        pdf.ln(1) # Line break

                # 4. Byte Output taiyar karein
                pdf_bytes = bytes(pdf.output())

            # 5. Download Button
            st.success("✅ Conversion Successful!")
            st.download_button(
                label="📥 Download PDF",
                data=pdf_bytes,
                file_name="word_converted.pdf",
                mime="application/pdf"
            )
            
        except Exception as e:
            st.error(f"Error: {e}")
# =====================================================
# TOOL: IMAGE -> PDF (NEW FEATURE)
# =====================================================
elif tool == "Image → PDF":
    st.header("🖼️ Convert Images to PDF")
    
    # Multiple images upload karne ka option
    img_files = st.file_uploader("Upload Images (JPG/PNG)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    
    if img_files and st.button("Convert to PDF"):
        try:
            with st.spinner("Combining images..."):
                # 1. Saari images ko open karke RGB mode me convert karein
                pil_images = []
                for img_file in img_files:
                    img = Image.open(img_file).convert("RGB")
                    pil_images.append(img)
                
                # 2. PDF me save karein
                if pil_images:
                    buf = io.BytesIO()
                    # Pehli image ko save karte hain aur baki ko append karte hain
                    pil_images[0].save(buf, format="PDF", save_all=True, append_images=pil_images[1:])
                    
                    pdf_bytes = buf.getvalue()

                    st.success(f"✅ Success! {len(pil_images)} images combined.")
                    
                    # 3. Download Button
                    st.download_button(
                        label="📥 Download PDF",
                        data=pdf_bytes,
                        file_name="images_combined.pdf",
                        mime="application/pdf"
                    )
                else:
                    st.error("No valid images found.")
                    
        except Exception as e:
            st.error(f"Error: {e}")
            

# =====================================================
# SECURITY
